"""
docx_builder.py  —  v8
======================================
FIXES in v8 vs v7:

1. DYNAMIC INDEX TERMS:
   Previously hardcoded as "RAG, FAISS, Knowledge Graph, LLM, Semantic Citation,
   IEEE Paper Generation, NLP" for every paper regardless of topic.
   Now computed from the actual topic + section content using TF-IDF-style
   frequency analysis. No hardcoded domain words anywhere.

2. BLOCK DIAGRAM DUPLICATION FIXED:
   _write_block_diagram() was rendering the diagram content twice because
   re.split() returns [before, content, after] and both the before-text and the
   after-text could contain a text repetition of the diagram. Added a
   diagram_rendered flag so the table is only drawn once.

3. All layout fixes from v7 preserved (Abstract+References in 2-column,
   graph sizing, algorithm blocks, etc.).
"""

import io
import re
import base64
from collections import Counter
from typing import List, Dict, Optional, Tuple

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Column widths ──────────────────────────────────────────────────────────────
_SINGLE_COL_IMG_W = Inches(5.8)
_DOUBLE_COL_IMG_W = Inches(3.2)

_FIGURE_COUNTER = [0]


def _next_figure() -> int:
    _FIGURE_COUNTER[0] += 1
    return _FIGURE_COUNTER[0]


# ═══════════════════════════════════════════════════════════════════════════════
# DYNAMIC INDEX TERMS  (replaces hardcoded string)
# ═══════════════════════════════════════════════════════════════════════════════

def _generate_index_terms(topic: str, sections: Dict[str, str]) -> str:
    """
    Derive IEEE Index Terms from the topic and section content.

    Uses token frequency across the abstract + introduction (most signal-dense
    sections) with a stopword filter. Falls back to the topic words if the
    sections are empty. Returns a comma-separated string of 5-7 terms.

    This replaces the previous hardcoded string:
    "RAG, FAISS, Knowledge Graph, LLM, Semantic Citation, IEEE Paper Generation, NLP"
    which was wrong for every non-paper-generator topic.
    """
    # Generic English stopwords + academic filler — no domain words in this list
    stopwords = {
        "the", "a", "an", "of", "in", "for", "and", "or", "to", "is", "are",
        "this", "that", "with", "by", "be", "as", "from", "on", "at", "it",
        "its", "we", "our", "their", "which", "these", "those", "has", "have",
        "been", "was", "were", "will", "can", "may", "also", "such", "each",
        "than", "into", "over", "under", "using", "used", "use", "used",
        "show", "shows", "shown", "present", "presents", "paper", "study",
        "work", "works", "approach", "method", "methods", "proposed", "based",
        "system", "systems", "model", "models", "data", "result", "results",
        "performance", "evaluation", "experiment", "experiments", "however",
        "while", "where", "when", "then", "thus", "therefore", "between",
        "both", "well", "further", "novel", "new", "existing", "current",
        "recent", "first", "second", "third", "high", "low", "large", "small",
        "different", "various", "several", "many", "number", "set", "sets",
        "order", "section", "table", "figure", "equation", "ieee", "journal",
    }

    # Use abstract and introduction — they have the most concentrated domain terms
    source_text = " ".join([
        topic,
        sections.get("abstract", ""),
        sections.get("introduction", "")[:800],
        sections.get("methodology", "")[:400],
    ])

    # Extract tokens: 4+ chars, letters only, lowercase
    tokens = re.findall(r'[a-zA-Z]{4,}', source_text)
    tokens_lower = [t.lower() for t in tokens if t.lower() not in stopwords]

    freq = Counter(tokens_lower)

    # Score: frequency × length bonus (longer technical terms are more specific)
    scored = [
        (word, count * (1 + min(len(word) - 4, 6) * 0.1))
        for word, count in freq.items()
        if count >= 1 and len(word) >= 4
    ]
    scored.sort(key=lambda x: -x[1])

    # Title-case the top terms; deduplicate by stemmed prefix
    seen_prefixes = set()
    terms = []
    for word, _ in scored:
        prefix = word[:5]  # rough dedup: "detection" and "detecting" → same prefix
        if prefix not in seen_prefixes:
            seen_prefixes.add(prefix)
            terms.append(word.title())
        if len(terms) >= 7:
            break

    # If we got very little (short abstract), fall back to topic words
    if len(terms) < 3:
        topic_words = [w.title() for w in re.findall(r'[a-zA-Z]{4,}', topic)
                       if w.lower() not in stopwords]
        terms = list(dict.fromkeys(terms + topic_words))[:7]

    return ", ".join(terms[:7]) if terms else topic


# ═══════════════════════════════════════════════════════════════════════════════
# TEXT CLEANING
# ═══════════════════════════════════════════════════════════════════════════════

_INLINE_REF_START = re.compile(
    r'\n\s*(?:\*\*References?\*\*|References?|REFERENCES?)\s*\n',
    re.IGNORECASE,
)
_MD_TABLE_SEP = re.compile(r'^\|?\s*[-:]+\s*\|')


def _clean_section_text(text: str, section_key: str) -> str:
    if not text:
        return text

    if section_key == "abstract":
        text = re.sub(r'(?i)^\s*abstract\s*[:\-–—]?\s*\n?', '', text.strip())

    m = _INLINE_REF_START.search(text)
    if m:
        text = text[:m.start()].rstrip()

    section_label_map = {
        "introduction":      r'(?i)^\s*\*{0,2}introduction\*{0,2}\s*$',
        "literature_survey": r'(?i)^\s*\*{0,2}literature\s+survey.*\*{0,2}\s*$',
        "methodology":       r'(?i)^\s*\*{0,2}methodology\*{0,2}\s*$',
        "algorithms":        r'(?i)^\s*\*{0,2}algorithms?\s+section\*{0,2}\s*$',
        "block_diagram":     r'(?i)^\s*\*{0,2}(system\s+)?architecture.*\*{0,2}\s*$',
        "results":           r'(?i)^\s*\*{0,2}results(\s+and\s+discussion)?\*{0,2}\s*$',
        "conclusion":        r'(?i)^\s*\*{0,2}conclusions?\*{0,2}\s*$',
    }
    if section_key in section_label_map:
        lines = text.splitlines()
        cleaned, skipped = [], False
        for line in lines:
            if not skipped and re.match(section_label_map[section_key], line):
                skipped = True
                continue
            cleaned.append(line)
        text = "\n".join(cleaned)

    return text.strip()


# ═══════════════════════════════════════════════════════════════════════════════
# LOW-LEVEL HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def _run(para, text, name="Times New Roman", size=10,
         bold=False, italic=False, color=None, superscript=False):
    run = para.add_run(text)
    run.font.name   = name
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    if superscript:
        rPr = run._element.get_or_add_rPr()
        va  = OxmlElement("w:vertAlign")
        va.set(qn("w:val"), "superscript")
        rPr.append(va)
    return run


def _para(doc, text="", align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=10,
          bold=False, italic=False, before=0, after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after  = Pt(after)
    if text:
        _run(p, text, size=size, bold=bold, italic=italic)
    return p


def _heading(doc, text, size=10):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    _run(p, text, size=size, bold=True)
    return p


def _rule(doc):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "AAAAAA")
    pBdr.append(bot)
    pPr.append(pBdr)


def _add_para_border(para, top=False, bottom=False):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for side, enabled in [("top", top), ("bottom", bottom)]:
        if enabled:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "6")
            el.set(qn("w:space"), "1")
            el.set(qn("w:color"), "2E4057")
            pBdr.append(el)
    pPr.append(pBdr)


# ═══════════════════════════════════════════════════════════════════════════════
# COLUMN LAYOUT
# ═══════════════════════════════════════════════════════════════════════════════

def _set_section_columns(section, num_cols: int, space_twips: int = 500):
    sectPr = section._sectPr
    for existing in sectPr.findall(qn("w:cols")):
        sectPr.remove(existing)
    cols = OxmlElement("w:cols")
    cols.set(qn("w:num"),        str(num_cols))
    cols.set(qn("w:space"),      str(space_twips))
    cols.set(qn("w:equalWidth"), "1")
    sectPr.append(cols)


# ═══════════════════════════════════════════════════════════════════════════════
# INLINE MARKDOWN
# ═══════════════════════════════════════════════════════════════════════════════

def _apply_inline_bold(para, text: str, size=10):
    parts = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|\[\d+\])', text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            _run(para, part[2:-2], size=size, bold=True)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            _run(para, part[1:-1], size=size, italic=True)
        elif re.fullmatch(r'\[\d+\]', part):
            _run(para, part[1:-1], size=8, superscript=True, color=(0, 0, 180))
        else:
            _run(para, part, size=size)


# ═══════════════════════════════════════════════════════════════════════════════
# MARKDOWN TABLE
# ═══════════════════════════════════════════════════════════════════════════════

def _is_md_table_block(lines: List[str]) -> bool:
    pipe_lines = [l for l in lines if l.strip().startswith("|") or "|" in l]
    return len(pipe_lines) >= 2


def _render_md_table(doc, lines: List[str]):
    rows_raw = []
    for line in lines:
        line = line.strip()
        if not line or _MD_TABLE_SEP.match(line):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows_raw.append(cells)

    if not rows_raw:
        return

    ncols    = max(len(r) for r in rows_raw)
    rows_raw = [r + [""] * (ncols - len(r)) for r in rows_raw]

    tbl = doc.add_table(rows=len(rows_raw), cols=ncols)
    tbl.style = "Table Grid"

    col_w = int(Inches(3.0) / ncols)
    for ri, row_data in enumerate(rows_raw):
        for ci, cell_text in enumerate(row_data):
            cell = tbl.rows[ri].cells[ci]
            cell.width = col_w
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(p, cell_text, size=8, bold=(ri == 0))
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# CODE FENCE
# ═══════════════════════════════════════════════════════════════════════════════

def _render_code_block(doc, code: str, lang: str = ""):
    if not code.strip():
        return
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(4)
    h.paragraph_format.space_after  = Pt(0)
    _add_para_border(h, top=True)
    r = h.add_run(lang.upper() if lang else "CODE")
    r.font.name = "Courier New"; r.font.size = Pt(7); r.font.bold = True
    r.font.color.rgb = RGBColor(46, 64, 87)

    for line in code.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.left_indent  = Inches(0.15)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        r2 = p.add_run(line if line else " ")
        r2.font.name = "Courier New"; r2.font.size = Pt(7.5)
        r2.font.color.rgb = RGBColor(20, 20, 20)

    foot = doc.add_paragraph()
    foot.paragraph_format.space_before = Pt(0)
    foot.paragraph_format.space_after  = Pt(6)
    _add_para_border(foot, bottom=True)


# ═══════════════════════════════════════════════════════════════════════════════
# PARAGRAPH WRITER
# ═══════════════════════════════════════════════════════════════════════════════

def _write_markdown_para(doc, text: str, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    stripped = text.strip()
    if not stripped:
        return

    if stripped.startswith("###"):
        _heading(doc, stripped.lstrip("#").strip(), size=9)
        return
    if stripped.startswith("##"):
        _heading(doc, stripped.lstrip("#").strip(), size=10)
        return
    if re.fullmatch(r'\*\*[^*]+\*\*', stripped):
        _heading(doc, stripped[2:-2], size=10)
        return

    bm = re.match(r'^[\*\-]\s+(.+)', stripped)
    if bm:
        try:
            p = doc.add_paragraph(style="List Bullet")
        except Exception:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(2)
        _apply_inline_bold(p, bm.group(1), size=10)
        return

    nm = re.match(r'^\d+[\.\)]\s+(.+)', stripped)
    if nm:
        try:
            p = doc.add_paragraph(style="List Number")
        except Exception:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(2)
        _apply_inline_bold(p, nm.group(1), size=10)
        return

    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(5)
    _apply_inline_bold(p, stripped, size=10)


def _write_body_text(doc, text: str, graph_map: Dict, in_two_col: bool = True):
    fence_re  = re.compile(r'```(\w*)\n(.*?)```', re.DOTALL)
    figure_re = re.compile(r'\[FIGURE_(\d+)\]', re.IGNORECASE)

    pos = 0
    for m in fence_re.finditer(text):
        _write_text_chunk(doc, text[pos:m.start()], graph_map, figure_re, in_two_col)
        _render_code_block(doc, m.group(2), lang=m.group(1))
        pos = m.end()
    _write_text_chunk(doc, text[pos:], graph_map, figure_re, in_two_col)


def _write_text_chunk(doc, text: str, graph_map: Dict, figure_re, in_two_col: bool):
    lines = text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        if "|" in line and (line.strip().startswith("|") or line.strip().count("|") >= 2):
            table_lines = []
            while i < len(lines) and "|" in lines[i]:
                table_lines.append(lines[i])
                i += 1
            if _is_md_table_block(table_lines):
                _render_md_table(doc, table_lines)
            else:
                for tl in table_lines:
                    _write_line_with_figures(doc, tl, graph_map, figure_re, in_two_col)
            continue

        _write_line_with_figures(doc, line, graph_map, figure_re, in_two_col)
        i += 1


def _write_line_with_figures(doc, line: str, graph_map: Dict, figure_re, in_two_col: bool):
    parts = figure_re.split(line)
    text_parts = []
    for j, part in enumerate(parts):
        if j % 2 == 0:
            text_parts.append(part)
        else:
            combined = " ".join(text_parts).strip()
            if combined:
                _write_markdown_para(doc, combined)
            text_parts = []
            key = f"fig_{part.strip()}"
            if key in graph_map:
                _embed_graph(doc, graph_map[key], _next_figure(), in_two_col)

    combined = " ".join(text_parts).strip()
    if combined:
        _write_markdown_para(doc, combined)


# ═══════════════════════════════════════════════════════════════════════════════
# ALGORITHM BLOCK
# ═══════════════════════════════════════════════════════════════════════════════

def _normalise_algo_text(text: str) -> str:
    text = re.sub(r'```\w*\n?', '', text)
    text = re.sub(r'```', '', text)
    text = re.sub(r'(?m)^#{1,3}\s*(ALGORITHM)', r'\1', text, flags=re.IGNORECASE)
    text = re.sub(r'\*{1,2}(ALGORITHM[^*]*)\*{1,2}', r'\1', text, flags=re.IGNORECASE)
    return text


def _parse_algorithm_blocks(text: str) -> List[Tuple[str, str]]:
    text   = _normalise_algo_text(text)
    blocks = []
    delim  = re.compile(r'(ALGORITHM[\s\S]*?END\s+ALGORITHM)', re.IGNORECASE)
    for m in delim.finditer(text):
        lines = [l.rstrip() for l in m.group(1).splitlines() if l.strip()]
        if lines:
            blocks.append((lines[0].strip(), "\n".join(lines[1:])))
    if not blocks:
        parts = re.split(r'(?im)(?=^ALGORITHM\s+)', text)
        for part in parts:
            part = part.strip()
            if not re.match(r'(?i)ALGORITHM\s+', part):
                continue
            lines = [l.rstrip() for l in part.splitlines() if l.strip()]
            blocks.append((lines[0].strip(), "\n".join(lines[1:])))
    return blocks


def _write_algorithm_block(doc, text: str):
    if not text or not text.strip():
        _para(doc, "No algorithm content.", size=9)
        return

    blocks    = _parse_algorithm_blocks(text)
    normalised = _normalise_algo_text(text)
    first_algo = re.search(r'(?i)ALGORITHM\s+', normalised)
    prefix = normalised[:first_algo.start()].strip() if first_algo else ""

    if prefix:
        for line in prefix.split("\n"):
            if line.strip():
                _write_markdown_para(doc, line.strip())

    if not blocks:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        r = p.add_run(normalised[:2000])
        r.font.name = "Courier New"; r.font.size = Pt(8)
        return

    for header, body in blocks:
        h_p = doc.add_paragraph()
        h_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        h_p.paragraph_format.space_before = Pt(6)
        h_p.paragraph_format.space_after  = Pt(0)
        _add_para_border(h_p, top=True)
        r = h_p.add_run(header)
        r.font.name = "Courier New"; r.font.size = Pt(8.5); r.font.bold = True
        r.font.color.rgb = RGBColor(46, 64, 87)

        body_lines = [l for l in body.splitlines() if l.strip()]
        for idx, line in enumerate(body_lines):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Inches(0.3)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after  = Pt(1)
            is_end = bool(re.match(r'(?i)END\s+ALGORITHM', line.strip()))
            if idx == len(body_lines) - 1:
                _add_para_border(p, bottom=True)
            r2 = p.add_run(line)
            r2.font.name = "Courier New"; r2.font.size = Pt(8.5)
            r2.font.bold = is_end


# ═══════════════════════════════════════════════════════════════════════════════
# BLOCK DIAGRAM  —  FIXED: no longer renders twice
# ═══════════════════════════════════════════════════════════════════════════════

def _write_block_diagram(doc, text: str):
    bd_re = re.compile(
        r'\[BLOCK_DIAGRAM_START\](.*?)\[BLOCK_DIAGRAM_END\]',
        re.DOTALL | re.IGNORECASE,
    )
    parts = bd_re.split(text)

    # re.split on a single-group pattern returns [before, group1, after, group2, ...]
    # We render the diagram table ONCE (first occurrence) then only write prose parts.
    diagram_rendered = False

    for i, part in enumerate(parts):
        if not part.strip():
            continue
        if i % 2 == 1:
            # This is a captured diagram block
            if not diagram_rendered:
                _render_block_diagram_table(doc, part.strip())
                diagram_rendered = True
            # If diagram_rendered is already True (duplicate tags in LLM output),
            # we skip re-rendering — the content would be identical anyway.
        else:
            # This is surrounding prose — write it, but skip any lines that
            # just repeat the [BLOCK_DIAGRAM_START/END] markers (LLM sometimes
            # includes them as text too)
            for line in part.split("\n"):
                stripped = line.strip()
                if not stripped:
                    continue
                if re.match(r'\[BLOCK_DIAGRAM_(START|END)\]', stripped, re.IGNORECASE):
                    continue  # skip stray marker lines
                _write_markdown_para(doc, stripped)


def _render_block_diagram_table(doc, diagram_text: str):
    lines       = [l.strip() for l in diagram_text.split("\n") if l.strip()]
    system_name = "System Architecture"
    components  = []
    submodules  = {}
    in_comps = in_subs = False

    for line in lines:
        if line.upper().startswith("SYSTEM:"):
            system_name = line.split(":", 1)[1].strip()
        elif line.upper().rstrip(":") in ("COMPONENTS", "SUBMODULES"):
            in_comps = line.upper().rstrip(":") == "COMPONENTS"
            in_subs  = not in_comps
        elif in_comps and ("->" in line or "→" in line):
            components.append(line)
        elif in_subs and ":" in line:
            k, v = line.split(":", 1)
            submodules[k.strip()] = v.strip()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(6)
    title_p.paragraph_format.space_after  = Pt(3)
    _run(title_p, f"Fig. — {system_name}", size=9, bold=True, color=(46, 64, 87))

    if not components:
        _para(doc, diagram_text, size=9)
        return

    tbl = doc.add_table(rows=len(components), cols=3)
    tbl.style = "Table Grid"
    for ri, comp_line in enumerate(components):
        row = tbl.rows[ri]
        m   = re.match(r'(.+?)\s*(?:->|→)\s*(.+?)(?::\s*(.*))?$', comp_line)
        src, dst, desc = (
            (m.group(1).strip(), m.group(2).strip(), (m.group(3) or "").strip())
            if m else (comp_line, "", "")
        )
        c0 = row.cells[0]
        c0.paragraphs[0].clear()
        c0.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = c0.paragraphs[0].add_run(src)
        r0.font.name = "Times New Roman"; r0.font.size = Pt(8); r0.font.bold = True
        r0.font.color.rgb = RGBColor(46, 64, 87)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D6E4F0")
        c0._tc.get_or_add_tcPr().append(shd)
        c1 = row.cells[1]
        c1.paragraphs[0].clear()
        c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        ar = c1.paragraphs[0].add_run("→")
        ar.font.size = Pt(11); ar.font.bold = True
        ar.font.color.rgb = RGBColor(46, 64, 87)
        c2 = row.cells[2]
        c2.paragraphs[0].clear()
        rd = c2.paragraphs[0].add_run(dst)
        rd.font.name = "Times New Roman"; rd.font.size = Pt(8); rd.font.bold = True
        if desc:
            ri2 = c2.paragraphs[0].add_run(f"  — {desc}")
            ri2.font.name = "Times New Roman"; ri2.font.size = Pt(7.5)
            ri2.font.italic = True; ri2.font.color.rgb = RGBColor(80, 80, 80)
        sub = submodules.get(dst, "")
        if sub:
            p2 = c2.add_paragraph()
            rs = p2.add_run(f"  [{sub}]")
            rs.font.name = "Times New Roman"; rs.font.size = Pt(7.5)
            rs.font.color.rgb = RGBColor(100, 100, 100)

    for row in tbl.rows:
        row.cells[0].width = Inches(1.4)
        row.cells[1].width = Inches(0.3)
        row.cells[2].width = Inches(3.0)
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# GRAPH EMBEDDING
# ═══════════════════════════════════════════════════════════════════════════════

def _validate_b64(data: str) -> bool:
    if not data or not isinstance(data, str):
        return False
    try:
        return len(base64.b64decode(data)) > 100
    except Exception:
        return False


def _embed_graph(doc, graph: Dict, figure_num: int, in_two_col: bool = True):
    if not _validate_b64(graph.get("data", "")):
        print(f"[docx_builder] Skipping '{graph.get('title')}' — invalid image")
        return

    img_width = _DOUBLE_COL_IMG_W if in_two_col else _SINGLE_COL_IMG_W

    img_p = doc.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_p.paragraph_format.space_before = Pt(8)
    img_p.paragraph_format.space_after  = Pt(2)
    try:
        img_p.add_run().add_picture(
            io.BytesIO(base64.b64decode(graph["data"])),
            width=img_width
        )
    except Exception as e:
        print(f"[docx_builder] Embed failed for '{graph.get('title')}': {e}")
        return

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after  = Pt(4)
    _run(cap, f"Fig. {figure_num}. ", size=8, bold=True)
    _run(cap, graph.get("title", ""), size=8, italic=True)

    insight = (graph.get("project_insight") or
               graph.get("statistical_insight") or
               graph.get("insight", ""))
    if insight:
        ins = doc.add_paragraph()
        ins.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        ins.paragraph_format.space_before = Pt(1)
        ins.paragraph_format.space_after  = Pt(8)
        ins.paragraph_format.left_indent  = Inches(0.2)
        ins.paragraph_format.right_indent = Inches(0.2)
        _run(ins, f"Analysis: {insight[:200]}", size=8, italic=True, color=(70, 70, 70))


# ═══════════════════════════════════════════════════════════════════════════════
# IEEE REFERENCE FORMATTER
# ═══════════════════════════════════════════════════════════════════════════════

def _format_ieee_reference(ref_num: int, citation: Dict) -> str:
    text = citation.get("text", "")
    if text.startswith(f"[{ref_num}]"):
        return text
    title   = citation.get("title", "")
    authors = citation.get("authors", "")
    year    = citation.get("year", "")
    journal = citation.get("journal", "")
    doi     = citation.get("doi", "")

    def abbrev(name: str) -> str:
        p = name.strip().split()
        return f"{p[0][0]}. {' '.join(p[1:])}" if len(p) >= 2 else name

    if authors:
        alist      = [abbrev(a.strip()) for a in re.split(r',\s*|\s+and\s+', authors)[:3]]
        author_str = ", ".join(alist)
    else:
        author_str = "Unknown"

    ref = f'[{ref_num}] {author_str}, "{title},"'
    if journal:
        ref += f" {journal},"
    if year:
        ref += f" {year}."
    if doi:
        ref += f" doi: {doi}."
    return ref if ref.strip() != f'[{ref_num}] Unknown, ""' else text


# ═══════════════════════════════════════════════════════════════════════════════
# METRICS PAGE
# ═══════════════════════════════════════════════════════════════════════════════

def _write_metrics_page(doc, metrics: Dict):
    if not metrics:
        return

    _heading(doc, "Appendix A: Paper Quality Metrics", size=10)

    data = [
        ("Metric", "Value", "Interpretation"),
        ("Perplexity", str(metrics.get("perplexity", "—")), "< 50 = coherent"),
        ("BLEU Score", str(metrics.get("bleu_score", "—")), "0–1; higher is better vs reference"),
        ("ROUGE-1 F1", str(metrics.get("rouge_1_f1", "—")), "Unigram overlap"),
        ("ROUGE-2 F1", str(metrics.get("rouge_2_f1", "—")), "Bigram overlap"),
        ("ROUGE-L F1", str(metrics.get("rouge_l_f1", "—")), "Longest common subsequence"),
        ("Lexical Diversity", str(metrics.get("lexical_diversity", "—")), "Type-token ratio; > 0.5 = varied"),
        ("Avg Sentence Length", str(metrics.get("avg_sentence_length", "—")), "Words per sentence"),
    ]

    tbl = doc.add_table(rows=len(data), cols=3)
    tbl.style = "Table Grid"
    for ri, row_data in enumerate(data):
        for ci, cell_text in enumerate(row_data):
            cell = tbl.rows[ri].cells[ci]
            p    = cell.paragraphs[0]
            _run(p, cell_text, size=9, bold=(ri == 0))
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN BUILDER  —  v8
# ═══════════════════════════════════════════════════════════════════════════════

def build_ieee_docx(
    output_path: str,
    topic: str,
    sections: Dict[str, str],
    citations: List[Dict],
    graphs: List[Dict],
    dataset_title: str = "",
    metrics: Optional[Dict] = None,
) -> str:
    """
    Build IEEE two-column DOCX.

    v8 CHANGES vs v7:
      - Index Terms are now dynamically derived from topic + section content
        (no more hardcoded "RAG, FAISS, NLP..." for every paper)
      - Block diagram duplication fixed (diagram_rendered flag)
      - All v7 layout preserved: Abstract + References in 2-column
    """
    _FIGURE_COUNTER[0] = 0

    doc       = Document()
    graph_map = {g["id"]: g for g in graphs}

    # ── Page margins ──────────────────────────────────────────────────────────
    sec0 = doc.sections[0]
    sec0.page_width    = Cm(21.59)
    sec0.page_height   = Cm(27.94)
    sec0.left_margin   = sec0.right_margin  = Cm(1.9)
    sec0.top_margin    = sec0.bottom_margin = Cm(2.2)
    _set_section_columns(sec0, 1)

    # ── Title ─────────────────────────────────────────────────────────────────
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp.paragraph_format.space_after  = Pt(6)
    tp.paragraph_format.space_before = Pt(0)
    _run(tp, f"A Comprehensive Study on {topic}",
         name="Times New Roman", size=18, bold=True)

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_after = Pt(2)
    _run(sp, "IEEE Format Research Paper — AI-Powered Automated Generation",
         size=10, italic=True)

    if dataset_title:
        dp = doc.add_paragraph()
        dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dp.paragraph_format.space_after = Pt(4)
        _run(dp, f"Dataset: {dataset_title}", size=9, italic=True, color=(60, 60, 100))

    _rule(doc)

    # ── SECTION BREAK → Switch to TWO-COLUMN ─────────────────────────────────
    body_sec = doc.add_section(WD_SECTION.CONTINUOUS)
    body_sec.left_margin   = sec0.left_margin
    body_sec.right_margin  = sec0.right_margin
    body_sec.top_margin    = sec0.top_margin
    body_sec.bottom_margin = sec0.bottom_margin
    _set_section_columns(body_sec, 2, space_twips=450)

    # ── ABSTRACT ──────────────────────────────────────────────────────────────
    ah = doc.add_paragraph()
    ah.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ah.paragraph_format.space_before = Pt(4)
    ah.paragraph_format.space_after  = Pt(3)
    _run(ah, "Abstract—", size=10, bold=True)

    abs_text = _clean_section_text(sections.get("abstract", ""), "abstract")
    abs_p = doc.add_paragraph()
    abs_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    abs_p.paragraph_format.space_after = Pt(4)
    _apply_inline_bold(abs_p, abs_text.replace("\n", " ").strip(), size=9)

    # ── INDEX TERMS — dynamically derived, never hardcoded ───────────────────
    kw_p = doc.add_paragraph()
    kw_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    kw_p.paragraph_format.space_after = Pt(8)
    _run(kw_p, "Index Terms—", size=9, italic=True, bold=True)
    index_terms = _generate_index_terms(topic, sections)
    _run(kw_p, index_terms, size=9, italic=True)

    # ── Body sections I – VII ─────────────────────────────────────────────────
    section_headings = [
        ("introduction",      "I. INTRODUCTION"),
        ("literature_survey", "II. LITERATURE SURVEY"),
        ("methodology",       "III. METHODOLOGY"),
        ("algorithms",        "IV. ALGORITHMS"),
        ("block_diagram",     "V. SYSTEM ARCHITECTURE"),
        ("results",           "VI. RESULTS AND DISCUSSION"),
        ("conclusion",        "VII. CONCLUSION"),
    ]

    for key, heading_text in section_headings:
        _heading(doc, heading_text, size=10)
        raw  = sections.get(key, "")
        text = _clean_section_text(raw, key)

        if key == "algorithms":
            _write_algorithm_block(doc, text)
        elif key == "block_diagram":
            _write_block_diagram(doc, text)
        elif key == "results":
            has_fig_tags = bool(re.search(r'\[FIGURE_\d+\]', text, re.IGNORECASE))
            _write_body_text(doc, text, graph_map, in_two_col=True)
            if not has_fig_tags and graphs:
                for g in graphs:
                    _embed_graph(doc, g, _next_figure(), in_two_col=True)
        else:
            _write_body_text(doc, text, graph_map, in_two_col=True)

    # ── REFERENCES ────────────────────────────────────────────────────────────
    rh = doc.add_paragraph()
    rh.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rh.paragraph_format.space_before = Pt(10)
    rh.paragraph_format.space_after  = Pt(4)
    _run(rh, "REFERENCES", size=10, bold=True)

    for i, cit in enumerate(citations, 1):
        ref_text = _format_ieee_reference(i, cit)
        rp = doc.add_paragraph()
        rp.paragraph_format.left_indent       = Inches(0.3)
        rp.paragraph_format.first_line_indent = Inches(-0.3)
        rp.paragraph_format.space_after       = Pt(3)
        _run(rp, ref_text, size=8.5)

    # ── METRICS APPENDIX ──────────────────────────────────────────────────────
    if metrics:
        _write_metrics_page(doc, metrics)

    doc.save(output_path)
    print(f"[docx_builder v8] Saved → {output_path} "
          f"({_FIGURE_COUNTER[0]} figures, {len(citations)} refs)")
    return output_path